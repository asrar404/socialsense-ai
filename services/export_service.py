import csv
import io
import json
import os
from datetime import datetime, timezone
from flask import current_app
from database import db
from models.report_export import ReportExport
from repositories.comment_result_repository import CommentResultRepository
from repositories.analysis_repository import AnalysisRepository
from repositories.report_export_repository import ReportExportRepository
from services.v12_context_service import (
    build_v12_context,
    HEURISTIC_DISCLAIMER,
    THREAT_DISCLAIMER,
)


_UNAVAILABLE = 'Unavailable'
_INSUFFICIENT = 'Insufficient evidence'


def _display(value, default=_UNAVAILABLE):
    """Render a score/field for exports, never showing a missing value as 0."""
    if value is None or (isinstance(value, (list, dict)) and not value):
        return default
    if isinstance(value, float) and value != value:  # NaN
        return default
    return value


def _fmt_number(value):
    if value is None:
        return _UNAVAILABLE
    try:
        return f'{float(value):.2f}'
    except (TypeError, ValueError):
        return _display(value)


def _fmt_pct(value):
    if value is None:
        return _UNAVAILABLE
    try:
        return f'{float(value) * 100:.0f}%'
    except (TypeError, ValueError):
        return _display(value)


def _evidence_text(value):
    """Render an evidence entry as a single display string.

    Narrative evidence samples are stored as small dicts
    ``{'source', 'ref', 'snippet'}``; the snippet is what we surface. Raw
    dicts/list values must never reach the binary writers (openpyxl raises on
    non-primitive cell values)."""
    if isinstance(value, dict):
        snippet = value.get('snippet')
        return str(snippet) if snippet is not None else _UNAVAILABLE
    if isinstance(value, (list, tuple)):
        return ' | '.join(str(_evidence_text(v)) for v in value)
    return _display(value)


class ExportService:
    def __init__(self):
        self.comment_repo = CommentResultRepository()
        self.analysis_repo = AnalysisRepository()
        self.export_repo = ReportExportRepository()

    def _build_export_bundle(self, analysis_id, user_id):
        """Assemble the full V1-V12 export data contract.

        Reuses existing read services/data only; no intelligence is recomputed
        here. Returns None if the analysis is missing/not owned by the user.
        """
        analysis = self.analysis_repo.get_user_analysis_with_reddit(
            analysis_id, user_id)
        if not analysis:
            return None

        comments = self.comment_repo.get_by_analysis_id(analysis_id)
        yt = analysis.youtube_analysis
        reddit = analysis.reddit_analysis

        from models.media_analysis import MediaAnalysis
        from models.entity import Entity
        from models.entity_context import EntityContext

        media = MediaAnalysis.query.filter_by(analysis_id=analysis_id).first()
        entities = Entity.query.filter_by(analysis_id=analysis_id).order_by(
            Entity.importance_score.desc()).limit(20).all()
        entity_data = []
        if entities:
            from services.entity_summary_service import EntitySummaryService
            eids = [e.id for e in entities]
            sentiments = EntityContext.query.filter(
                EntityContext.entity_id.in_(eids)).with_entities(
                EntityContext.entity_id, EntityContext.entity_sentiment,
                EntityContext.entity_sentiment_score,
            ).all()
            risks = EntityContext.query.filter(
                EntityContext.entity_id.in_(eids)).with_entities(
                EntityContext.entity_id, EntityContext.entity_risk_score,
            ).all()
            sent_group = {}
            for eid, sent, score in sentiments:
                sent_group.setdefault(eid, []).append(
                    {'sentiment': sent, 'score': score})
            risk_group = {}
            for eid, score in risks:
                risk_group.setdefault(eid, []).append(score)
            sent_summary = []
            risk_summary = []
            for e in entities:
                avg_s = round(sum(s['score'] for s in sent_group.get(e.id, []))
                              / max(len(sent_group.get(e.id, [])), 1), 1)
                overall = 'positive' if avg_s > 60 else \
                    'negative' if avg_s < 40 else 'neutral'
                sent_summary.append({'entity_name': e.normalized_name,
                                     'overall_sentiment': overall,
                                     'average_score': avg_s})
                avg_r = round(sum(risk_group.get(e.id, [0]))
                              / max(len(risk_group.get(e.id, [])), 1), 1)
                risk_summary.append({'entity_name': e.normalized_name,
                                     'average_risk_score': avg_r})
            for e, s, r in zip(entities, sent_summary, risk_summary):
                entity_data.append({
                    'name': e.name,
                    'normalized_name': e.normalized_name,
                    'entity_type': e.entity_type,
                    'frequency': e.frequency,
                    'importance_score': e.importance_score,
                    'overall_sentiment': s['overall_sentiment'],
                    'average_sentiment_score': s['average_score'],
                    'average_risk_score': r['average_risk_score'],
                })

        v12 = build_v12_context(analysis_id, user_id)

        return {
            'analysis': analysis,
            'youtube': yt,
            'reddit': reddit,
            'comments': comments,
            'media': media,
            'entities': entity_data,
            'v12': v12,
            'heuristic_disclaimer': HEURISTIC_DISCLAIMER,
            'threat_disclaimer': THREAT_DISCLAIMER,
        }

    def _export_record(self, analysis_id, format_type, filepath):
        record = ReportExport(
            analysis_id=analysis_id,
            format_type=format_type,
            file_path=filepath,
        )
        db.session.add(record)
        db.session.commit()

    def _save_file(self, filename, mode='w', encoding='utf-8'):
        filepath = os.path.join(
            current_app.config.get('UPLOAD_FOLDER', 'reports'), filename)
        os.makedirs(os.path.dirname(filepath), exist_ok=True)
        return filepath

    @staticmethod
    def _write_csv_rows(writer, rows):
        for row in rows:
            writer.writerow([_display(c) if c is None else c for c in row])

    def _write_v12_csv(self, writer, v12):
        """Append bounded V12 intelligence sections to a CSV writer."""
        threat = v12.get('threat')
        if threat:
            writer.writerow(['# Threat Assessment (heuristic, non-causal)'])
            writer.writerow(['Overall Threat Score',
                             _fmt_number(threat.get('overall_threat_score'))])
            writer.writerow(['Threat Level', _display(threat.get('threat_level'))])
            writer.writerow(['Confidence',
                             _fmt_number(threat.get('confidence'))])
            writer.writerow(['Evidence Coverage',
                             _fmt_pct(threat.get('evidence_coverage'))])
            comps = threat.get('component_scores') or {}
            for key in ('authenticity', 'coordination', 'narrative',
                        'propagation', 'temporal', 'entity'):
                label = key.replace('_', ' ').title()
                writer.writerow([f'Component: {label}',
                                 _display(comps.get(key))])
            writer.writerow(['Summary', _display(threat.get('summary'))])
            for reason in (threat.get('reasons') or [])[:12]:
                writer.writerow(['Reason', reason])
            for limitation in (threat.get('limitations') or [])[:8]:
                writer.writerow(['Limitation', limitation])
            writer.writerow([])

        narratives = v12.get('narratives') or []
        if narratives:
            writer.writerow(['# Narratives & Growth (heuristic, observed)'])
            for n in narratives[:8]:
                writer.writerow(['Narrative', _display(n.get('name'))])
                writer.writerow(['Category', _display(n.get('category'))])
                writer.writerow(['Risk Level', _display(n.get('risk_level'))])
                writer.writerow(['Risk Score', _fmt_number(n.get('risk_score'))])
                writer.writerow(['Growth Score', _fmt_number(n.get('growth_score'))])
                writer.writerow(['Confidence', _fmt_number(n.get('confidence'))])
                writer.writerow(['Occurrences', _display(n.get('occurrence_count'))])
                writer.writerow(['Cross-Platform',
                                 'Yes' if n.get('is_cross_platform') else 'No'])
                for kw in (n.get('keywords') or [])[:8]:
                    writer.writerow(['Keyword', kw])
            writer.writerow([])

        coordination = v12.get('coordination') or []
        if coordination:
            writer.writerow(['# Coordination Signals (heuristic, observed)'])
            for s in coordination[:8]:
                writer.writerow(['Signal Type', _display(s.get('signal_type'))])
                writer.writerow(['Level', _display(s.get('level'))])
                writer.writerow(['Score', _fmt_number(s.get('score'))])
                writer.writerow(['Confidence', _fmt_number(s.get('confidence'))])
                if s.get('cluster_size'):
                    writer.writerow(['Cluster Size', s.get('cluster_size')])
                if s.get('comparisons_performed') is not None:
                    writer.writerow(['Comparisons Performed',
                                     s.get('comparisons_performed')])
                    writer.writerow(['Comparisons Truncated',
                                     'Yes' if s.get('comparisons_truncated') else 'No'])
                for reason in (s.get('reasons') or [])[:6]:
                    writer.writerow(['Signal Reason', reason])
            writer.writerow([])

        propagation = v12.get('propagation') or []
        if propagation:
            writer.writerow(['# Propagation Relationships (heuristic, observed)'])
            for e in propagation[:8]:
                writer.writerow(['Source Platform', _display(e.get('source_platform'))])
                writer.writerow(['Target Platform', _display(e.get('target_platform'))])
                writer.writerow(['Relationship Type', _display(e.get('relationship_type'))])
                writer.writerow(['Propagation Score', _fmt_number(e.get('propagation_score'))])
                writer.writerow(['Similarity Score', _fmt_number(e.get('similarity_score'))])
                if e.get('lag_seconds') is not None:
                    writer.writerow(['Lag (hours)',
                                     f"{float(e['lag_seconds']) / 3600:.1f}"])
                for reason in (e.get('reasons') or [])[:6]:
                    writer.writerow(['Propagation Reason', reason])
            writer.writerow([])

        temporal = v12.get('temporal') or []
        if temporal:
            writer.writerow(['# Temporal Intelligence (heuristic, observed)'])
            for t in temporal[:8]:
                writer.writerow(['Narrative', _display(t.get('normalized_name'))])
                writer.writerow(['Growth Score', _fmt_number(t.get('growth_score'))])
                for key, value in (t.get('signals') or {}).items():
                    if value is not None:
                        writer.writerow([f'Temporal {key.replace("_", " ").title()}',
                                         value])
            writer.writerow([])

    def _write_evidence_csv(self, writer, bundle):
        """Append bounded Evidence sections (V11 media + V12 reasons)."""
        media = bundle.get('media')
        if media:
            writer.writerow(['# Evidence: Authenticity (heuristic)'])
            if media.reasons:
                try:
                    reasons = json.loads(media.reasons)
                except (TypeError, ValueError):
                    reasons = []
                for reason in (reasons or [])[:10]:
                    writer.writerow(['Evidence Reason', reason])
            writer.writerow([])

        v12 = bundle.get('v12') or {}
        rows = []
        for n in (v12.get('narratives') or [])[:8]:
            evidence = n.get('evidence') or {}
            samples = (evidence.get('samples') or [])[:3]
            for sample in samples:
                rows.append(('Narrative Evidence', n.get('name'), sample))
        for s in (v12.get('coordination') or [])[:8]:
            for reason in (s.get('reasons') or [])[:6]:
                rows.append(('Coordination Evidence', s.get('signal_type'), reason))
        for e in (v12.get('propagation') or [])[:8]:
            for reason in (e.get('reasons') or [])[:6]:
                rows.append(('Propagation Evidence',
                             f"{e.get('source_platform')}->{e.get('target_platform')}",
                             reason))
        if rows:
            writer.writerow(['# Evidence: V12 Signals (heuristic, non-causal)'])
            writer.writerow(['Source', 'Label', 'Detail'])
            for source, label, detail in rows[:30]:
                writer.writerow([source, label, detail])
            writer.writerow([])

    def generate_csv(self, analysis_id, user_id):
        bundle = self._build_export_bundle(analysis_id, user_id)
        if not bundle:
            return None

        analysis = bundle['analysis']
        comments = bundle['comments']
        yt = bundle['youtube']
        reddit = bundle['reddit']

        output = io.StringIO()
        writer = csv.writer(output)

        if analysis.analysis_type == 'reddit' and reddit:
            writer.writerow(['# Reddit Post Metadata'])
            writer.writerow(['Post ID', reddit.post_id])
            writer.writerow(['Subreddit', reddit.subreddit or ''])
            writer.writerow(['Title', reddit.post_title or ''])
            writer.writerow(['Author', reddit.post_author or ''])
            writer.writerow(['Score', reddit.post_score or 0])
            writer.writerow(['Upvote Ratio', reddit.upvote_ratio or 0.0])
            writer.writerow(['Comment Limit', reddit.comment_limit or 100])
            writer.writerow(['Demo Mode', 'Yes' if reddit.is_demo else 'No'])
        else:
            writer.writerow(['# Video Metadata'])
            writer.writerow(['Video ID', yt.video_id if yt else ''])
            writer.writerow(['Title', yt.video_title if yt else ''])
            writer.writerow(['Channel', yt.channel_name if yt else ''])
            writer.writerow(['Views', yt.view_count if yt else 0])
            writer.writerow(['Likes', yt.like_count if yt else 0])
            writer.writerow(['Comment Limit', yt.comment_limit if yt else 100])
            writer.writerow(['Demo Mode', 'Yes' if (yt and yt.is_demo) else 'No'])
        writer.writerow([])

        if analysis.analysis_summary:
            writer.writerow(['# Analysis Summary'])
            for line in analysis.analysis_summary.split('\n'):
                writer.writerow([line])
            writer.writerow([])

        from models.media_analysis import MediaAnalysis
        media = MediaAnalysis.query.filter_by(analysis_id=analysis_id).first()
        if media:
            writer.writerow(['# Authenticity Intelligence (heuristic)'])
            writer.writerow(['Overall AI Probability', media.overall_ai_probability])
            writer.writerow(['Overall Authenticity Score', media.overall_authenticity_score])
            writer.writerow(['Confidence', media.confidence])
            writer.writerow(['Deepfake Score', media.deepfake_score])
            writer.writerow(['Synthetic Voice Score', media.synthetic_voice_score])
            writer.writerow(['Thumbnail AI Score', media.thumbnail_ai_score])
            writer.writerow(['Frame Manipulation Score', media.frame_manipulation_score])
            writer.writerow(['Metadata Score', media.metadata_score])
            writer.writerow(['Summary', media.summary or ''])
            if media.reasons:
                try:
                    reasons = json.loads(media.reasons)
                    for reason in reasons:
                        writer.writerow(['Reason', reason])
                except (ValueError, TypeError):
                    writer.writerow(['Reason', media.reasons])
            writer.writerow([])

        from models.video_context_history import VideoContextHistory
        vch = VideoContextHistory.query.filter_by(analysis_id=analysis_id).first()
        history_context_score = vch.avg_sentiment if vch else ''
        history_risk_score = vch.avg_risk if vch else ''

        writer.writerow([
            'Comment',
            'Author',
            'Published At',
            'Sentiment',
            'Sentiment Score',
            'Sentiment Confidence',
            'Sentiment Explanation',
            'Spam Score',
            'Spam Confidence',
            'Spam Explanation',
            'Toxicity Score',
            'Toxicity Confidence',
            'Toxicity Explanation',
            'Duplicate Score',
            'Duplicate Explanation',
            'Bot Score',
            'Bot Confidence',
            'Bot Explanation',
            'Risk Score',
            'Risk Level',
            'Risk Explanation',
            'Recommendation',
            'Context Relevance Score',
            'Context Match Label',
            'Context Reason',
            'Entities',
            'Entity Types',
            'Entity Sentiments',
            'Entity Risk Scores',
            'Entity Relevance Scores',
            'Historical Context Score',
            'Historical Risk Score',
            'Entity Recurrence Score',
            'Topic Recurrence Score',
            'Trend Direction',
        ])

        for c in comments:
            published = c.published_at.strftime('%Y-%m-%d %H:%M:%S UTC') if c.published_at else ''
            ctx = c.context
            writer.writerow([
                c.comment_text,
                c.author or '',
                published,
                c.sentiment or '',
                c.sentiment_score,
                c.sentiment_confidence or '',
                c.sentiment_explanation or '',
                c.spam_score,
                c.spam_confidence or '',
                c.spam_explanation or '',
                c.toxicity_score,
                c.toxicity_confidence or '',
                c.toxicity_explanation or '',
                c.duplicate_score,
                c.duplicate_explanation or '',
                c.bot_score,
                c.bot_confidence or '',
                c.bot_explanation or '',
                c.risk_score,
                c.risk_level,
                c.risk_explanation or '',
                c.recommendation or '',
                ctx.transcript_relevance_score if ctx else '',
                ctx.context_match_label.replace('_', ' ').title() if ctx else '',
                ctx.reason if ctx else '',
                '; '.join([em.entity.name for em in c.entity_mentions.all()]) if c.entity_mentions.count() > 0 else '',
                '; '.join([em.entity.entity_type for em in c.entity_mentions.all()]) if c.entity_mentions.count() > 0 else '',
                '; '.join([ec.entity_sentiment or '' for ec in c.entity_contexts.all()]) if c.entity_contexts.count() > 0 else '',
                '; '.join([str(ec.entity_risk_score) for ec in c.entity_contexts.all()]) if c.entity_contexts.count() > 0 else '',
                '; '.join([str(ec.entity_relevance_score) for ec in c.entity_contexts.all()]) if c.entity_contexts.count() > 0 else '',
                history_context_score,
                history_risk_score,
                '',
                '',
                '',
            ])

        writer.writerow([])
        self._write_v12_csv(writer, bundle['v12'])
        self._write_evidence_csv(writer, bundle)

        csv_content = output.getvalue()
        output.close()

        filename = f"analysis_{analysis_id}_{datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S')}.csv"
        filepath = os.path.join(current_app.config.get('UPLOAD_FOLDER', 'reports'), filename)

        os.makedirs(os.path.dirname(filepath), exist_ok=True)
        with open(filepath, 'w', newline='', encoding='utf-8') as f:
            f.write(csv_content)

        export_record = ReportExport(
            analysis_id=analysis_id,
            format_type='csv',
            file_path=filepath,
        )
        db.session.add(export_record)
        db.session.commit()

        return {
            'csv_content': csv_content,
            'filename': filename,
            'filepath': filepath,
        }

    def generate_json(self, analysis_id, user_id):
        analysis = self.analysis_repo.get_user_analysis_with_reddit(analysis_id, user_id)
        if not analysis:
            return None

        comments = self.comment_repo.get_by_analysis_id(analysis_id)
        yt = analysis.youtube_analysis
        reddit = analysis.reddit_analysis

        metadata = {}
        if analysis.analysis_type == 'reddit' and reddit:
            metadata = {
                'platform': 'reddit',
                'post_id': reddit.post_id,
                'subreddit': reddit.subreddit,
                'title': reddit.post_title,
                'body': reddit.post_body,
                'author': reddit.post_author,
                'score': reddit.post_score,
                'upvote_ratio': reddit.upvote_ratio,
                'comment_limit': reddit.comment_limit,
                'permalink': reddit.permalink,
                'created_utc': reddit.created_utc.isoformat() if reddit.created_utc else None,
                'is_demo': reddit.is_demo,
            }
        elif yt:
            metadata = {
                'platform': 'youtube',
                'video_id': yt.video_id,
                'title': yt.video_title,
                'description': yt.video_description,
                'channel': yt.channel_name,
                'published_at': yt.published_at.isoformat() if yt.published_at else None,
                'view_count': yt.view_count,
                'like_count': yt.like_count,
                'comment_limit': yt.comment_limit,
                'is_demo': yt.is_demo,
            }

        comments_data = []
        for c in comments:
            ctx = c.context
            comments_data.append({
                'comment_text': c.comment_text,
                'author': c.author,
                'published_at': c.published_at.isoformat() if c.published_at else None,
                'sentiment': c.sentiment,
                'sentiment_score': c.sentiment_score,
                'sentiment_confidence': c.sentiment_confidence,
                'sentiment_explanation': c.sentiment_explanation,
                'spam_score': c.spam_score,
                'spam_confidence': c.spam_confidence,
                'spam_explanation': c.spam_explanation,
                'toxicity_score': c.toxicity_score,
                'toxicity_confidence': c.toxicity_confidence,
                'toxicity_explanation': c.toxicity_explanation,
                'duplicate_score': c.duplicate_score,
                'duplicate_explanation': c.duplicate_explanation,
                'ai_like_score': c.ai_like_score,
                'ai_like_explanation': c.ai_like_explanation,
                'bot_score': c.bot_score,
                'bot_confidence': c.bot_confidence,
                'bot_explanation': c.bot_explanation,
                'risk_score': c.risk_score,
                'risk_level': c.risk_level,
                'risk_explanation': c.risk_explanation,
                'recommendation': c.recommendation,
                'context_relevance_score': ctx.transcript_relevance_score if ctx else None,
                'context_match_label': ctx.context_match_label.replace('_', ' ').title() if ctx else None,
                'context_reason': ctx.reason if ctx else None,
                'entity_mentions': [{'name': em.entity.name, 'type': em.entity.entity_type} for em in c.entity_mentions.all()] if c.entity_mentions.count() > 0 else [],
                'entity_sentiments': [{'entity': ec.entity.name if ec.entity else '', 'sentiment': ec.entity_sentiment, 'score': ec.entity_sentiment_score} for ec in c.entity_contexts.all()] if c.entity_contexts.count() > 0 else [],
                'entity_risks': [{'entity': ec.entity.name if ec.entity else '', 'risk_score': ec.entity_risk_score} for ec in c.entity_contexts.all()] if c.entity_contexts.count() > 0 else [],
            })

        data = {
            'exported_at': datetime.now(timezone.utc).isoformat(),
            'analysis_id': analysis_id,
            'analysis_type': analysis.analysis_type,
            'created_at': analysis.created_at.isoformat() if analysis.created_at else None,
            'metadata': metadata,
            'analysis_summary': analysis.analysis_summary,
            'comment_count': len(comments_data),
            'comments': comments_data,
        }

        from models.media_analysis import MediaAnalysis
        media = MediaAnalysis.query.filter_by(analysis_id=analysis_id).first()
        if media:
            data['media_analysis'] = media.to_dict()

        from models.entity import Entity
        from services.entity_summary_service import EntitySummaryService
        entity_list = Entity.query.filter_by(analysis_id=analysis_id).all()
        if entity_list:
            data['entity_summary'] = EntitySummaryService().generate_summary(entity_list, [], [])

        from models.video_context_history import VideoContextHistory
        from models.entity_history import EntityHistory
        from models.channel_context import ChannelContext
        vch = VideoContextHistory.query.filter_by(analysis_id=analysis_id).first()
        if vch:
            data['video_context_history'] = {
                'video_id': vch.video_id,
                'channel_id': vch.channel_id,
                'entity_count': vch.entity_count,
                'avg_sentiment': vch.avg_sentiment,
                'avg_risk': vch.avg_risk,
                'top_entities': vch.top_entities,
            }
        ehs = EntityHistory.query.filter_by(analysis_id=analysis_id).all()
        if ehs:
            data['entity_history'] = [
                {
                    'normalized_name': eh.normalized_name,
                    'entity_type': eh.entity_type,
                    'sentiment_score': eh.sentiment_score,
                    'risk_score': eh.risk_score,
                    'mention_count': eh.mention_count,
                }
                for eh in ehs
            ]

        v12 = build_v12_context(analysis_id, user_id)
        data['v12'] = {
            'threat': v12.get('threat'),
            'narratives': v12.get('narratives') or [],
            'coordination': v12.get('coordination') or [],
            'propagation': v12.get('propagation') or [],
            'temporal': v12.get('temporal') or [],
            'heuristic_disclaimer': HEURISTIC_DISCLAIMER,
            'threat_disclaimer': THREAT_DISCLAIMER,
        }

        from models.entity import Entity
        entity_list = Entity.query.filter_by(analysis_id=analysis_id).order_by(
            Entity.importance_score.desc()).limit(20).all()
        if entity_list:
            data['entities'] = [
                {
                    'name': e.name,
                    'normalized_name': e.normalized_name,
                    'entity_type': e.entity_type,
                    'frequency': e.frequency,
                    'importance_score': e.importance_score,
                }
                for e in entity_list
            ]

        json_content = json.dumps(data, indent=2, default=str)

        filename = f"analysis_{analysis_id}_{datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S')}.json"
        filepath = os.path.join(current_app.config.get('UPLOAD_FOLDER', 'reports'), filename)

        os.makedirs(os.path.dirname(filepath), exist_ok=True)
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(json_content)

        export_record = ReportExport(
            analysis_id=analysis_id,
            format_type='json',
            file_path=filepath,
        )
        db.session.add(export_record)
        db.session.commit()

        return {
            'json_content': json_content,
            'filename': filename,
            'filepath': filepath,
        }


    # ------------------------------------------------------------------ shared
    # Structured section model shared by XLSX / PDF / DOCX. Each section is a
    # dict: {title, subtitle, disclaimer, rows (list of [label, value])}.

    def _metadata_rows(self, bundle):
        analysis = bundle['analysis']
        yt = bundle['youtube']
        reddit = bundle['reddit']
        rows = []
        if analysis.analysis_type == 'reddit' and reddit:
            rows += [
                ['Platform', 'Reddit'],
                ['Post ID', _display(reddit.post_id)],
                ['Subreddit', _display(reddit.subreddit)],
                ['Title', _display(reddit.post_title)],
                ['Author', _display(reddit.post_author)],
                ['Score', _display(reddit.post_score)],
                ['Upvote Ratio', _display(reddit.upvote_ratio)],
                ['Comment Limit', _display(reddit.comment_limit)],
                ['Demo Mode', 'Yes' if reddit.is_demo else 'No'],
            ]
        else:
            rows += [
                ['Platform', 'YouTube'],
                ['Video ID', _display(yt.video_id if yt else None)],
                ['Title', _display(yt.video_title if yt else None)],
                ['Channel', _display(yt.channel_name if yt else None)],
                ['Views', _display(yt.view_count if yt else None)],
                ['Likes', _display(yt.like_count if yt else None)],
                ['Comment Limit', _display(yt.comment_limit if yt else None)],
                ['Demo Mode', 'Yes' if (yt and yt.is_demo) else 'No'],
            ]
        rows += [
            ['Analysis Type', analysis.analysis_type],
            ['Created At', analysis.created_at.isoformat()
             if analysis.created_at else _UNAVAILABLE],
            ['Export Time', datetime.now(timezone.utc).isoformat()],
        ]
        if analysis.analysis_summary:
            rows.append(['Summary', analysis.analysis_summary])
        return rows

    def _threat_section(self, bundle):
        threat = (bundle.get('v12') or {}).get('threat')
        rows = []
        if not threat:
            return {'title': 'Threat Assessment',
                    'subtitle': None,
                    'disclaimer': THREAT_DISCLAIMER,
                    'rows': [['Overall Threat Score', _INSUFFICIENT]],
                    'unavailable': True}
        comps = threat.get('component_scores') or {}
        rows += [
            ['Overall Threat Score', _fmt_number(threat.get('overall_threat_score'))],
            ['Threat Level', _display(threat.get('threat_level'))],
            ['Confidence', _fmt_number(threat.get('confidence'))],
            ['Evidence Coverage', _fmt_pct(threat.get('evidence_coverage'))],
            ['Assessment Method', _display(threat.get('assessment_method'))],
        ]
        for key in ('authenticity', 'coordination', 'narrative',
                    'propagation', 'temporal', 'entity'):
            rows.append([f'Component: {key.replace("_", " ").title()}',
                         _display(comps.get(key))])
        caps = threat.get('capability_labels') or {}
        if caps:
            for key, label in caps.items():
                rows.append([f'Capability: {key.replace("_", " ").title()}',
                             _display(label)])
        if threat.get('summary'):
            rows.append(['Summary', threat.get('summary')])
        for reason in (threat.get('reasons') or [])[:12]:
            rows.append(['Reason', reason])
        for limitation in (threat.get('limitations') or [])[:8]:
            rows.append(['Limitation', limitation])
        return {'title': 'Threat Assessment',
                'subtitle': 'Heuristic weighted assessment',
                'disclaimer': THREAT_DISCLAIMER,
                'rows': rows,
                'unavailable': False}

    def _authenticity_section(self, bundle):
        media = bundle.get('media')
        rows = []
        if not media:
            return {'title': 'Authenticity Intelligence',
                    'subtitle': 'Heuristic assessment',
                    'disclaimer': HEURISTIC_DISCLAIMER,
                    'rows': [['Overall AI Probability', _INSUFFICIENT]],
                    'unavailable': True}
        rows += [
            ['Overall AI Probability', _fmt_number(media.overall_ai_probability)],
            ['Overall Authenticity Score',
             _fmt_number(media.overall_authenticity_score)],
            ['Confidence', _fmt_number(media.confidence)],
            ['Deepfake Score', _fmt_number(media.deepfake_score)],
            ['Synthetic Voice Score', _fmt_number(media.synthetic_voice_score)],
            ['Thumbnail AI Score', _fmt_number(media.thumbnail_ai_score)],
            ['Frame Manipulation Score', _fmt_number(media.frame_manipulation_score)],
            ['Metadata Score', _fmt_number(media.metadata_score)],
        ]
        if media.summary:
            rows.append(['Summary', media.summary])
        if media.reasons:
            try:
                reasons = json.loads(media.reasons)
            except (TypeError, ValueError):
                reasons = []
            for reason in (reasons or [])[:10]:
                rows.append(['Reason', reason])
        return {'title': 'Authenticity Intelligence',
                'subtitle': 'Heuristic, explainable assessment',
                'disclaimer': HEURISTIC_DISCLAIMER,
                'rows': rows,
                'unavailable': False}

    def _narrative_section(self, bundle):
        narratives = (bundle.get('v12') or {}).get('narratives') or []
        rows = []
        for n in narratives[:8]:
            rows.append(['Narrative', _display(n.get('name'))])
            rows.append(['Category', _display(n.get('category'))])
            rows.append(['Risk Level', _display(n.get('risk_level'))])
            rows.append(['Risk Score', _fmt_number(n.get('risk_score'))])
            rows.append(['Growth Score', _fmt_number(n.get('growth_score'))])
            rows.append(['Confidence', _fmt_number(n.get('confidence'))])
            rows.append(['Occurrences', _display(n.get('occurrence_count'))])
            rows.append(['Cross-Platform',
                         'Yes' if n.get('is_cross_platform') else 'No'])
            for kw in (n.get('keywords') or [])[:8]:
                rows.append(['Keyword', kw])
        return {'title': 'Narratives & Growth',
                'subtitle': 'Heuristic, observed growth (not causal)',
                'disclaimer': HEURISTIC_DISCLAIMER,
                'rows': rows,
                'unavailable': not rows}

    def _coordination_section(self, bundle):
        coordination = (bundle.get('v12') or {}).get('coordination') or []
        rows = []
        for s in coordination[:8]:
            rows.append(['Signal Type', _display(s.get('signal_type'))])
            rows.append(['Level', _display(s.get('level'))])
            rows.append(['Score', _fmt_number(s.get('score'))])
            rows.append(['Confidence', _fmt_number(s.get('confidence'))])
            if s.get('cluster_size'):
                rows.append(['Cluster Size', s.get('cluster_size')])
            if s.get('comparisons_performed') is not None:
                rows.append(['Comparisons Performed',
                             s.get('comparisons_performed')])
                rows.append(['Comparisons Truncated',
                             'Yes' if s.get('comparisons_truncated') else 'No'])
            if s.get('summary'):
                rows.append(['Summary', s.get('summary')])
            for reason in (s.get('reasons') or [])[:6]:
                rows.append(['Reason', reason])
        return {'title': 'Coordination Signals',
                'subtitle': 'Heuristic, observed patterns (not proof)',
                'disclaimer': HEURISTIC_DISCLAIMER,
                'rows': rows,
                'unavailable': not rows}

    def _propagation_section(self, bundle):
        propagation = (bundle.get('v12') or {}).get('propagation') or []
        rows = []
        for e in propagation[:8]:
            rows.append(['Source Platform', _display(e.get('source_platform'))])
            rows.append(['Target Platform', _display(e.get('target_platform'))])
            rows.append(['Relationship Type', _display(e.get('relationship_type'))])
            rows.append(['Propagation Score', _fmt_number(e.get('propagation_score'))])
            rows.append(['Similarity Score', _fmt_number(e.get('similarity_score'))])
            rows.append(['Confidence', _fmt_number(e.get('confidence'))])
            if e.get('lag_seconds') is not None:
                rows.append(['Lag (hours)',
                             f"{float(e['lag_seconds']) / 3600:.1f}"])
            rows.append(['Cross-Platform',
                         'Yes' if e.get('is_cross_platform') else 'No'])
            for reason in (e.get('reasons') or [])[:6]:
                rows.append(['Reason', reason])
        return {'title': 'Propagation Relationships',
                'subtitle': 'Heuristic, observed relationships (not causal)',
                'disclaimer': HEURISTIC_DISCLAIMER,
                'rows': rows,
                'unavailable': not rows}

    def _temporal_section(self, bundle):
        temporal = (bundle.get('v12') or {}).get('temporal') or []
        rows = []
        for t in temporal[:8]:
            rows.append(['Narrative', _display(t.get('normalized_name'))])
            rows.append(['Growth Score', _fmt_number(t.get('growth_score'))])
            for key, value in (t.get('signals') or {}).items():
                if value is not None:
                    rows.append([f'Temporal {key.replace("_", " ").title()}',
                                 value])
            for reason in (t.get('reasons') or [])[:6]:
                rows.append(['Reason', reason])
        return {'title': 'Temporal Intelligence',
                'subtitle': 'Heuristic, observed growth over time',
                'disclaimer': HEURISTIC_DISCLAIMER,
                'rows': rows,
                'unavailable': not rows}

    def _entity_section(self, bundle):
        entities = bundle.get('entities') or []
        rows = []
        for e in entities[:20]:
            rows.append(['Entity', _display(e.get('name'))])
            rows.append(['Type', _display(e.get('entity_type'))])
            rows.append(['Frequency', _display(e.get('frequency'))])
            rows.append(['Importance Score', _fmt_number(e.get('importance_score'))])
            rows.append(['Overall Sentiment', _display(e.get('overall_sentiment'))])
            rows.append(['Avg Sentiment Score',
                         _fmt_number(e.get('average_sentiment_score'))])
            rows.append(['Avg Risk Score', _fmt_number(e.get('average_risk_score'))])
        return {'title': 'Entity Intelligence',
                'subtitle': 'Observed entity/context signals',
                'disclaimer': HEURISTIC_DISCLAIMER,
                'rows': rows,
                'unavailable': not rows}

    def _evidence_section(self, bundle):
        rows = []
        media = bundle.get('media')
        if media and media.reasons:
            try:
                reasons = json.loads(media.reasons)
            except (TypeError, ValueError):
                reasons = []
            for reason in (reasons or [])[:10]:
                rows.append(['Authenticity Evidence', reason])
        v12 = bundle.get('v12') or {}
        for n in (v12.get('narratives') or [])[:8]:
            evidence = n.get('evidence') or {}
            for sample in (evidence.get('samples') or [])[:3]:
                rows.append(['Narrative Evidence', _evidence_text(sample)])
        for s in (v12.get('coordination') or [])[:8]:
            for reason in (s.get('reasons') or [])[:6]:
                rows.append(['Coordination Evidence', _evidence_text(reason)])
        for e in (v12.get('propagation') or [])[:8]:
            for reason in (e.get('reasons') or [])[:6]:
                rows.append(['Propagation Evidence', _evidence_text(reason)])
        return {'title': 'Evidence',
                'subtitle': 'Bounded supporting evidence (heuristic, non-causal)',
                'disclaimer': HEURISTIC_DISCLAIMER,
                'rows': rows,
                'unavailable': not rows}

    def _all_sections(self, bundle):
        return [
            self._threat_section(bundle),
            self._authenticity_section(bundle),
            self._narrative_section(bundle),
            self._coordination_section(bundle),
            self._propagation_section(bundle),
            self._temporal_section(bundle),
            self._entity_section(bundle),
            self._evidence_section(bundle),
        ]

# ------------------------------------------------------------------ XLSX

    def generate_xlsx(self, analysis_id, user_id):
        bundle = self._build_export_bundle(analysis_id, user_id)
        if not bundle:
            return None

        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment

        wb = Workbook()
        warn_font = Font(italic=True, color='777777')

        def write_sheet(ws, section):
            ws.title = section['title'][:31]
            ws.cell(row=1, column=1, value=section['title']).font = Font(
                bold=True, size=13)
            if section.get('subtitle'):
                ws.cell(row=2, column=1, value=section['subtitle']).font = warn_font
            if section.get('disclaimer'):
                ws.cell(row=3, column=1,
                        value=section['disclaimer']).font = warn_font
            if section.get('unavailable') and not section['rows']:
                ws.cell(row=5, column=1, value=_INSUFFICIENT)
            row_idx = 5
            for label, value in section['rows']:
                ws.cell(row=row_idx, column=1, value=label).font = Font(bold=True)
                if isinstance(value, (dict, list, tuple)):
                    value = _evidence_text(value)
                cell = ws.cell(row=row_idx, column=2, value=value)
                cell.alignment = Alignment(wrap_text=True)
                row_idx += 1
            ws.column_dimensions['A'].width = 34
            ws.column_dimensions['B'].width = 70

        summary_ws = wb.active
        summary_ws.title = 'Summary'
        summary_ws.cell(row=1, column=1, value='Analysis Summary').font = Font(
            bold=True, size=13)
        summary_ws.cell(row=2, column=1,
                        value=HEURISTIC_DISCLAIMER).font = warn_font
        summary_ws.cell(row=4, column=1, value='Metadata').font = Font(bold=True)
        row_idx = 5
        for label, value in self._metadata_rows(bundle):
            summary_ws.cell(row=row_idx, column=1, value=label).font = Font(
                bold=True)
            cell = summary_ws.cell(row=row_idx, column=2, value=value)
            cell.alignment = Alignment(wrap_text=True)
            row_idx += 1
        summary_ws.column_dimensions['A'].width = 30
        summary_ws.column_dimensions['B'].width = 70

        for section in self._all_sections(bundle):
            write_sheet(wb.create_sheet(), section)

        filename = (f"analysis_{analysis_id}_"
                    f"{datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S')}.xlsx")
        filepath = self._save_file(filename, mode='wb')
        wb.save(filepath)

        self._export_record(analysis_id, 'xlsx', filepath)
        return {'filename': filename, 'filepath': filepath}

    # ------------------------------------------------------------------ PDF

    def generate_pdf(self, analysis_id, user_id):
        bundle = self._build_export_bundle(analysis_id, user_id)
        if not bundle:
            return None

        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import mm, cm
        from reportlab.lib import colors
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
            PageBreak, KeepTogether,
        )

        filename = (f"analysis_{analysis_id}_"
                    f"{datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S')}.pdf")
        filepath = self._save_file(filename, mode='wb')

        doc = SimpleDocTemplate(filepath, pagesize=A4,
                                leftMargin=20*mm, rightMargin=20*mm,
                                topMargin=20*mm, bottomMargin=20*mm)
        styles = getSampleStyleSheet()
        warn_style = ParagraphStyle(
            'WarnItalic', parent=styles['Normal'],
            fontSize=8, textColor=colors.HexColor('#777777'),
            spaceAfter=6)
        section_style = ParagraphStyle(
            'SectionTitle', parent=styles['Heading2'],
            spaceBefore=12, spaceAfter=4, fontSize=13, bold=True)
        subtitle_style = ParagraphStyle(
            'Subtitle', parent=warn_style, fontSize=9, spaceAfter=2)
        cell_style = ParagraphStyle(
            'Cell', parent=styles['Normal'], fontSize=8, leading=10)
        value_style = ParagraphStyle(
            'ValueCell', parent=styles['Normal'], fontSize=8, leading=10)

        elements = []
        elements.append(Paragraph('Analysis Report', styles['Title']))
        elements.append(Paragraph(HEURISTIC_DISCLAIMER, warn_style))
        elements.append(Spacer(1, 6*mm))
        elements.append(Paragraph('Metadata', styles['Heading3']))
        for label, value in self._metadata_rows(bundle):
            t = Table(
                [[Paragraph(f'<b>{label}</b>', cell_style),
                  Paragraph(str(value or ''), value_style)]],
                colWidths=[doc.width * 0.30, doc.width * 0.70])
            t.setStyle(TableStyle([
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('LEFTPADDING', (0, 0), (-1, -1), 2),
                ('RIGHTPADDING', (0, 0), (-1, -1), 2),
                ('TOPPADDING', (0, 0), (-1, -1), 1),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 1),
            ]))
            elements.append(t)

        for section in self._all_sections(bundle):
            elements.append(Spacer(1, 4*mm))
            elements.append(Paragraph(section['title'], section_style))
            if section.get('subtitle'):
                elements.append(Paragraph(section['subtitle'], subtitle_style))
            if section.get('disclaimer'):
                elements.append(Paragraph(section['disclaimer'], warn_style))
            if section.get('unavailable') and not section['rows']:
                elements.append(Paragraph(_INSUFFICIENT, styles['Normal']))
            else:
                for label, value in section['rows']:
                    t = Table(
                        [[Paragraph(f'<b>{label}</b>', cell_style),
                          Paragraph(str(value or ''), value_style)]],
                        colWidths=[doc.width * 0.30, doc.width * 0.70])
                    t.setStyle(TableStyle([
                        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                        ('LEFTPADDING', (0, 0), (-1, -1), 2),
                        ('RIGHTPADDING', (0, 0), (-1, -1), 2),
                        ('TOPPADDING', (0, 0), (-1, -1), 1),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 1),
                    ]))
                    elements.append(t)

        doc.build(elements)
        self._export_record(analysis_id, 'pdf', filepath)
        return {'filename': filename, 'filepath': filepath}

    # ------------------------------------------------------------------ DOCX

    def generate_docx(self, analysis_id, user_id):
        bundle = self._build_export_bundle(analysis_id, user_id)
        if not bundle:
            return None

        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        doc.add_heading('Analysis Report', 0)
        p = doc.add_paragraph()
        run = p.add_run(HEURISTIC_DISCLAIMER)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
        run.italic = True

        doc.add_heading('Metadata', level=2)
        for label, value in self._metadata_rows(bundle):
            p = doc.add_paragraph()
            r = p.add_run(f'{label}: ')
            r.bold = True
            p.add_run(str(value or ''))

        for section in self._all_sections(bundle):
            doc.add_heading(section['title'], level=2)
            if section.get('subtitle'):
                p = doc.add_paragraph()
                r = p.add_run(section['subtitle'])
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
                r.italic = True
            if section.get('disclaimer'):
                p = doc.add_paragraph()
                r = p.add_run(section['disclaimer'])
                r.font.size = Pt(8)
                r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
                r.italic = True
            if section.get('unavailable') and not section['rows']:
                doc.add_paragraph(_INSUFFICIENT)
            else:
                for label, value in section['rows']:
                    p = doc.add_paragraph()
                    r = p.add_run(f'{label}: ')
                    r.bold = True
                    p.add_run(str(value or ''))

        filename = (f"analysis_{analysis_id}_"
                    f"{datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S')}.docx")
        filepath = self._save_file(filename, mode='wb')
        doc.save(filepath)

        self._export_record(analysis_id, 'docx', filepath)
        return {'filename': filename, 'filepath': filepath}
